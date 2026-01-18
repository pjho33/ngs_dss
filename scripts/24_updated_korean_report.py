#!/usr/bin/env python3
"""
나이 효과를 고려한 업데이트된 한글 최종 리포트
"""

import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
warnings.filterwarnings('ignore')

BASE_DIR = Path("/home/pjho3/projects/ngs_dss")
RESULTS_DIR = BASE_DIR / "results"
OUTPUT_DIR = RESULTS_DIR / "final_korean_report"
OUTPUT_DIR.mkdir(exist_ok=True)

print("=" * 80)
print("나이 효과를 고려한 최종 리포트 생성")
print("=" * 80)

# Load results
age_adjusted = pd.read_csv(RESULTS_DIR / "age_adjusted_analysis" / "age_adjusted_summary.csv")
integrated_data = pd.read_csv(RESULTS_DIR / "complete_reanalysis" / "integrated_data_complete.csv")

patients = integrated_data[integrated_data['label'] == 1]
controls = integrated_data[integrated_data['label'] == 0]

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(11)

# Title
title = doc.add_heading('췌장암 조기 진단을 위한 EMR-seq 메틸화 분석 최종 보고서', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('(나이 효과 분석 포함)')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_para = doc.add_paragraph('분석 완료일: 2026년 1월 18일')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============================================================================
# 1. 연구 개요
# ============================================================================
doc.add_heading('1. 연구 개요', 1)

doc.add_heading('1.1 연구 목적', 2)
doc.add_paragraph(
    '본 연구는 EMR-seq (Enzymatic Methyl-seq) 기술을 이용하여 췌장암 환자와 건강한 대조군의 '
    'DNA 메틸화 패턴을 분석하고, 이를 기반으로 췌장암 조기 진단 모델을 개발하는 것을 목적으로 합니다.'
)

doc.add_heading('1.2 연구 대상', 2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = '구분'
hdr_cells[1].text = '환자군'
hdr_cells[2].text = '대조군'

table.rows[1].cells[0].text = '샘플 수'
table.rows[1].cells[1].text = str(len(patients))
table.rows[1].cells[2].text = str(len(controls))

table.rows[2].cells[0].text = '평균 나이'
table.rows[2].cells[1].text = f"{patients['age'].mean():.1f}세"
table.rows[2].cells[2].text = f"{controls['age'].mean():.1f}세"

table.rows[3].cells[0].text = 'CA19-9 평균'
table.rows[3].cells[1].text = f"{patients['ca19_9'].mean():.1f} U/mL"
table.rows[3].cells[2].text = f"{controls['ca19_9'].mean():.1f} U/mL"

doc.add_paragraph()

# ⚠️ 중요한 제한점 추가
warning_para = doc.add_paragraph()
warning_para.add_run('⚠️ 중요한 제한점: ').bold = True
warning_para.add_run(
    f'환자군과 대조군 간 평균 나이 차이가 {patients["age"].mean() - controls["age"].mean():.1f}세로 매우 큽니다. '
    '이는 분석 결과 해석에 중요한 영향을 미칩니다.'
)

doc.add_paragraph()

# 나이 분포
doc.add_heading('1.3 나이 분포', 2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = '나이대'
hdr_cells[1].text = '환자군'
hdr_cells[2].text = '대조군'

table.rows[1].cells[0].text = '20-39세'
table.rows[1].cells[1].text = f"{sum((patients['age'] >= 20) & (patients['age'] < 40))}명"
table.rows[1].cells[2].text = f"{sum((controls['age'] >= 20) & (controls['age'] < 40))}명"

table.rows[2].cells[0].text = '40-59세'
table.rows[2].cells[1].text = f"{sum((patients['age'] >= 40) & (patients['age'] < 60))}명"
table.rows[2].cells[2].text = f"{sum((controls['age'] >= 40) & (controls['age'] < 60))}명"

table.rows[3].cells[0].text = '60-79세'
table.rows[3].cells[1].text = f"{sum((patients['age'] >= 60) & (patients['age'] < 80))}명"
table.rows[3].cells[2].text = f"{sum((controls['age'] >= 60) & (controls['age'] < 80))}명"

table.rows[4].cells[0].text = '80세 이상'
table.rows[4].cells[1].text = f"{sum(patients['age'] >= 80)}명"
table.rows[4].cells[2].text = f"{sum(controls['age'] >= 80)}명"

doc.add_page_break()

# ============================================================================
# 2. 분석 방법
# ============================================================================
doc.add_heading('2. 분석 방법', 1)

doc.add_heading('2.1 메틸화 분석', 2)
doc.add_paragraph('• DMRseq 알고리즘을 사용하여 차별적으로 메틸화된 영역(DMR) 검출')
doc.add_paragraph('• 통계적 유의성: FDR < 0.05, 최소 3개 CpG sites, 최소 1000bp')
doc.add_paragraph('• 총 46개의 유의한 DMR 발견')

doc.add_heading('2.2 예측 모델', 2)
doc.add_paragraph('• Logistic Regression with balanced class weights')
doc.add_paragraph('• 5-fold stratified cross-validation')
doc.add_paragraph('• 여러 특징 조합으로 모델 성능 비교')

doc.add_heading('2.3 나이 효과 분석', 2)
doc.add_paragraph(
    '환자군과 대조군의 나이 차이가 크므로, 나이 효과를 분리하여 분석하였습니다:'
)
doc.add_paragraph('• 나이만 사용한 모델')
doc.add_paragraph('• 메틸화만 사용한 모델')
doc.add_paragraph('• 나이를 제외한 모델')
doc.add_paragraph('• 모든 변수를 포함한 모델')

doc.add_page_break()

# ============================================================================
# 3. 주요 결과
# ============================================================================
doc.add_heading('3. 주요 결과', 1)

doc.add_heading('3.1 나이 효과 분석', 2)

# Age effect table
table = doc.add_table(rows=len(age_adjusted)+1, cols=5)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = '모델'
hdr_cells[1].text = 'AUC'
hdr_cells[2].text = '정확도'
hdr_cells[3].text = '민감도'
hdr_cells[4].text = '특이도'

for idx, row in age_adjusted.iterrows():
    cells = table.rows[idx+1].cells
    cells[0].text = row['모델']
    cells[1].text = f"{row['AUC']:.3f}"
    cells[2].text = f"{row['정확도']:.3f}"
    cells[3].text = f"{row['민감도']:.3f}"
    cells[4].text = f"{row['특이도']:.3f}"

doc.add_paragraph()

# 핵심 발견
doc.add_heading('핵심 발견:', 3)

age_only = age_adjusted[age_adjusted['모델'] == '나이만'].iloc[0]
meth_only = age_adjusted[age_adjusted['모델'] == '메틸화만'].iloc[0]
no_age = age_adjusted[age_adjusted['모델'] == '메틸화 + 성별 + CA19-9 (나이 제외)'].iloc[0]
full_model = age_adjusted[age_adjusted['모델'] == '메틸화 + 나이 + 성별 + CA19-9 (원본)'].iloc[0]

warning_para = doc.add_paragraph()
warning_para.add_run('⚠️ 나이 효과가 매우 큽니다!\n').bold = True
warning_para.add_run(
    f'• 나이만으로 AUC {age_only["AUC"]:.3f}, 정확도 {age_only["정확도"]*100:.1f}% 달성\n'
    f'• 이는 모델이 실제로 "암"보다는 "나이"를 학습했을 가능성을 시사합니다.\n'
    f'• 환자군 평균 나이 {patients["age"].mean():.1f}세 vs 대조군 {controls["age"].mean():.1f}세'
)

doc.add_paragraph()

doc.add_heading('순수 메틸화 효과:', 3)
doc.add_paragraph(
    f'• 메틸화만 사용: AUC {meth_only["AUC"]:.3f}, 정확도 {meth_only["정확도"]*100:.1f}%'
)
doc.add_paragraph(
    f'• 나이 제외 (메틸화 + 성별 + CA19-9): AUC {no_age["AUC"]:.3f}, 정확도 {no_age["정확도"]*100:.1f}%'
)
doc.add_paragraph(
    '• 메틸화 데이터 자체는 진단 가치가 있으나, 나이 효과만큼 강력하지는 않습니다.'
)

doc.add_paragraph()

# Add age distribution image
doc.add_heading('나이 분포 및 ROC 곡선:', 3)
doc.add_picture(str(RESULTS_DIR / "age_adjusted_analysis" / "age_effect_analysis.png"), 
                width=Inches(6))

doc.add_page_break()

# ============================================================================
# 4. 연구의 제한점
# ============================================================================
doc.add_heading('4. 연구의 제한점', 1)

doc.add_heading('4.1 나이 불균형 문제', 2)
doc.add_paragraph(
    '본 연구의 가장 큰 제한점은 환자군과 대조군 간 나이 차이입니다:'
)

limitation_para = doc.add_paragraph()
limitation_para.add_run('주요 문제점:\n').bold = True
doc.add_paragraph(
    f'• 환자군 평균 나이: {patients["age"].mean():.1f}세 (대부분 60세 이상)'
)
doc.add_paragraph(
    f'• 대조군 평균 나이: {controls["age"].mean():.1f}세 (대부분 20-40대)'
)
doc.add_paragraph(
    f'• 평균 차이: {patients["age"].mean() - controls["age"].mean():.1f}세'
)
doc.add_paragraph(
    f'• 나이만으로 AUC {age_only["AUC"]:.3f} 달성 가능'
)

doc.add_paragraph()

doc.add_heading('4.2 해석상의 주의점', 2)
doc.add_paragraph(
    '현재 결과는 다음과 같이 해석되어야 합니다:'
)
doc.add_paragraph(
    '• 높은 전체 성능(AUC 0.978)은 주로 나이 차이에 기인합니다.'
)
doc.add_paragraph(
    '• 순수한 메틸화 효과는 AUC 0.912 수준입니다.'
)
doc.add_paragraph(
    '• 실제 임상 적용 시 같은 나이대에서의 성능은 더 낮을 수 있습니다.'
)

doc.add_heading('4.3 개선 방안', 2)
doc.add_paragraph('향후 연구에서 다음이 필요합니다:')
doc.add_paragraph('• 나이 매칭된 대조군 확보 (60세 이상)')
doc.add_paragraph('• 더 큰 규모의 나이 균형 잡힌 코호트')
doc.add_paragraph('• 나이를 보정한 통계 분석')
doc.add_paragraph('• 다기관 검증 연구')

doc.add_page_break()

# ============================================================================
# 5. 메틸화 데이터의 진단 가치
# ============================================================================
doc.add_heading('5. 메틸화 데이터의 진단 가치', 1)

doc.add_heading('5.1 순수 메틸화 효과', 2)
doc.add_paragraph(
    '나이 효과를 제외하고도 메틸화 데이터는 진단 가치가 있습니다:'
)
doc.add_paragraph(
    f'• 메틸화만 사용: AUC {meth_only["AUC"]:.3f}, 민감도 {meth_only["민감도"]*100:.1f}%'
)
doc.add_paragraph(
    '• 46개 DMR이 환자와 대조군을 유의하게 구분'
)
doc.add_paragraph(
    '• 비침습적 액체 생검으로 활용 가능'
)

doc.add_heading('5.2 CA19-9와의 비교', 2)
cancer_ca199 = patients['ca19_9'].dropna()
control_ca199 = controls['ca19_9'].dropna()

doc.add_paragraph(
    f'• CA19-9는 환자군에서 {cancer_ca199.mean() / control_ca199.mean():.1f}배 높음'
)
doc.add_paragraph(
    f'• 메틸화 데이터에 CA19-9 추가 시 AUC 소폭 향상 (+0.015)'
)
doc.add_paragraph(
    '• 메틸화 데이터가 CA19-9보다 더 포괄적인 정보 제공'
)

doc.add_page_break()

# ============================================================================
# 6. 결론
# ============================================================================
doc.add_heading('6. 결론', 1)

doc.add_heading('6.1 주요 발견', 2)
doc.add_paragraph(
    '본 연구는 EMR-seq 기반 메틸화 분석의 가능성과 제한점을 모두 보여줍니다:'
)

doc.add_paragraph()
doc.add_heading('긍정적 측면:', 3)
doc.add_paragraph(
    f'• 메틸화 데이터 자체는 진단 가치가 있음 (AUC {meth_only["AUC"]:.3f})'
)
doc.add_paragraph(
    '• 46개의 유의한 DMR 발견'
)
doc.add_paragraph(
    '• 비침습적 액체 생검 가능성'
)

doc.add_paragraph()
doc.add_heading('제한점:', 3)
doc.add_paragraph(
    f'• 나이 불균형으로 인한 편향 (나이만으로 AUC {age_only["AUC"]:.3f})'
)
doc.add_paragraph(
    '• 실제 임상 성능은 보고된 것보다 낮을 수 있음'
)
doc.add_paragraph(
    '• 나이 매칭된 대조군 필요'
)

doc.add_heading('6.2 임상적 의의', 2)
doc.add_paragraph(
    '나이 효과를 고려하더라도, 본 연구는 다음의 의의가 있습니다:'
)
doc.add_paragraph(
    '• 메틸화 기반 췌장암 진단의 가능성 입증'
)
doc.add_paragraph(
    '• 향후 연구 설계에 중요한 교훈 제공'
)
doc.add_paragraph(
    '• 나이 보정 분석의 중요성 강조'
)

doc.add_heading('6.3 향후 연구 방향', 2)
doc.add_paragraph('• 나이 매칭된 대규모 코호트 구축')
doc.add_paragraph('• 나이 보정 통계 모델 개발')
doc.add_paragraph('• 다기관 전향적 검증 연구')
doc.add_paragraph('• 다른 암종과의 감별 진단 연구')

doc.add_page_break()

# ============================================================================
# 7. 권장사항
# ============================================================================
doc.add_heading('7. 권장사항', 1)

doc.add_heading('7.1 결과 보고 시', 2)
doc.add_paragraph('본 연구 결과를 보고할 때는 반드시 다음을 명시해야 합니다:')
doc.add_paragraph(
    '• 환자군과 대조군 간 큰 나이 차이 (39.4세)'
)
doc.add_paragraph(
    f'• 나이만으로도 높은 성능 달성 가능 (AUC {age_only["AUC"]:.3f})'
)
doc.add_paragraph(
    f'• 순수 메틸화 효과는 AUC {meth_only["AUC"]:.3f} 수준'
)
doc.add_paragraph(
    '• 실제 임상 적용 시 성능은 더 낮을 수 있음'
)

doc.add_heading('7.2 추가 연구 필요성', 2)
doc.add_paragraph('다음 연구가 시급히 필요합니다:')
doc.add_paragraph(
    '• 60세 이상 건강한 대조군 최소 50명 이상 확보'
)
doc.add_paragraph(
    '• 나이 매칭된 케이스-컨트롤 연구 설계'
)
doc.add_paragraph(
    '• 나이를 공변량으로 하는 통계 분석'
)

doc.add_heading('7.3 임상 적용 시 고려사항', 2)
doc.add_paragraph(
    '현재 모델을 임상에 적용하기 전에:'
)
doc.add_paragraph(
    '• 같은 나이대 환자에서의 성능 재평가 필요'
)
doc.add_paragraph(
    '• 나이 보정 모델 개발 필요'
)
doc.add_paragraph(
    '• 독립적인 검증 코호트에서 성능 확인 필요'
)

doc.add_page_break()

# ============================================================================
# 8. 부록
# ============================================================================
doc.add_heading('8. 부록', 1)

doc.add_heading('8.1 통계 분석 상세', 2)
doc.add_paragraph('• DMR 검출: DMRseq (FDR < 0.05)')
doc.add_paragraph('• 예측 모델: Logistic Regression')
doc.add_paragraph('• 교차 검증: 5-fold stratified CV')
doc.add_paragraph('• 특징 정규화: StandardScaler')

doc.add_heading('8.2 데이터 요약', 2)
doc.add_paragraph(f'• 총 샘플: {len(integrated_data)}개')
doc.add_paragraph(f'• 환자군: {len(patients)}명 (평균 {patients["age"].mean():.1f}세)')
doc.add_paragraph(f'• 대조군: {len(controls)}명 (평균 {controls["age"].mean():.1f}세)')
doc.add_paragraph(f'• 유의한 DMR: 46개')

# Save
output_file = OUTPUT_DIR / "췌장암_EMRseq_분석_최종보고서_나이효과포함.docx"
doc.save(output_file)

print(f"\n업데이트된 한글 리포트 생성 완료!")
print(f"저장 위치: {output_file}")
print("\n주요 내용:")
print("  - 나이 효과 분석 포함")
print(f"  - 나이만으로 AUC {age_only['AUC']:.3f} 달성 명시")
print(f"  - 순수 메틸화 효과 AUC {meth_only['AUC']:.3f} 강조")
print("  - 연구 제한점 및 개선 방안 상세 기술")
print("=" * 80)
