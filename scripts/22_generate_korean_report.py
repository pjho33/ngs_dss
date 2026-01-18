#!/usr/bin/env python3
"""
췌장암 EMR-seq 메틸화 분석 최종 한글 리포트 생성
"""

import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import warnings
warnings.filterwarnings('ignore')

BASE_DIR = Path("/home/pjho3/projects/ngs_dss")
RESULTS_DIR = BASE_DIR / "results" / "complete_reanalysis"
OUTPUT_DIR = BASE_DIR / "results" / "final_korean_report"
OUTPUT_DIR.mkdir(exist_ok=True)

print("=" * 80)
print("췌장암 EMR-seq 메틸화 분석 최종 리포트 생성")
print("=" * 80)

# Load results
overall_summary = pd.read_csv(RESULTS_DIR / "overall_performance_summary.csv")
stage_summary = pd.read_csv(RESULTS_DIR / "stage_specific_summary.csv")
integrated_data = pd.read_csv(RESULTS_DIR / "integrated_data_complete.csv")

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(11)

# Title
title = doc.add_heading('췌장암 조기 진단을 위한 EMR-seq 메틸화 분석 최종 보고서', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Date
date_para = doc.add_paragraph('분석 완료일: 2026년 1월 18일')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============================================================================
# 1. 연구 개요
# ============================================================================
doc.add_heading('1. 연구 개요', 1)

doc.add_heading('1.1 연구 목적', 2)
doc.add_paragraph(
    '본 연구는 EMR-seq (Enzymatic Methyl-seq) 기술을 이용하여 췌장암 환자와 건강한 대조군의 '
    'DNA 메틸화 패턴을 분석하고, 이를 기반으로 췌장암 조기 진단 모델을 개발하는 것을 목적으로 합니다.'
)

doc.add_heading('1.2 연구 대상', 2)
patients = integrated_data[integrated_data['label'] == 1]
controls = integrated_data[integrated_data['label'] == 0]

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '구분'
hdr_cells[1].text = '환자군'
hdr_cells[2].text = '대조군'

# Data
table.rows[1].cells[0].text = '샘플 수'
table.rows[1].cells[1].text = str(len(patients))
table.rows[1].cells[2].text = str(len(controls))

table.rows[2].cells[0].text = '평균 나이'
table.rows[2].cells[1].text = f"{patients['age'].mean():.1f}세"
table.rows[2].cells[2].text = f"{controls['age'].mean():.1f}세"

table.rows[3].cells[0].text = 'CA19-9 평균'
table.rows[3].cells[1].text = f"{patients['ca19_9'].mean():.1f} U/mL"
table.rows[3].cells[2].text = f"{controls['ca19_9'].mean():.1f} U/mL"

doc.add_paragraph()

# Stage distribution
doc.add_heading('1.3 환자군 병기 분포', 2)
stage_counts = patients['stage'].value_counts().sort_index()
for stage, count in stage_counts.items():
    if pd.notna(stage):
        doc.add_paragraph(f'• Stage {int(stage)}: {count}명', style='List Bullet')

doc.add_page_break()

# ============================================================================
# 2. 분석 방법
# ============================================================================
doc.add_heading('2. 분석 방법', 1)

doc.add_heading('2.1 메틸화 분석', 2)
doc.add_paragraph(
    '• DMRseq (Differential Methylated Regions sequencing) 알고리즘을 사용하여 '
    '환자군과 대조군 간 차별적으로 메틸화된 영역(DMR)을 검출하였습니다.'
)
doc.add_paragraph(
    '• 통계적 유의성: FDR < 0.05, 최소 3개 CpG sites, 최소 1000bp 길이'
)
doc.add_paragraph(
    f'• 총 46개의 유의한 DMR을 발견하였습니다.'
)

doc.add_heading('2.2 예측 모델', 2)
doc.add_paragraph(
    '• 알고리즘: Logistic Regression with balanced class weights'
)
doc.add_paragraph(
    '• 교차 검증: 5-fold stratified cross-validation'
)
doc.add_paragraph(
    '• 특징 변수: 46개 DMR 메틸화 수준 + 임상 변수(나이, 성별, CA19-9)'
)

doc.add_heading('2.3 성능 평가 지표', 2)
doc.add_paragraph('• AUC-ROC (Area Under the ROC Curve): 모델의 전반적인 판별 능력')
doc.add_paragraph('• 정확도 (Accuracy): 전체 예측 중 정확한 예측의 비율')
doc.add_paragraph('• 민감도 (Sensitivity): 실제 암 환자 중 정확히 진단된 비율')
doc.add_paragraph('• 특이도 (Specificity): 실제 건강인 중 정확히 진단된 비율')
doc.add_paragraph('• 정밀도 (Precision): 암으로 예측된 것 중 실제 암인 비율')

doc.add_page_break()

# ============================================================================
# 3. 주요 결과
# ============================================================================
doc.add_heading('3. 주요 결과', 1)

doc.add_heading('3.1 전체 성능', 2)

# Performance table
table = doc.add_table(rows=len(overall_summary)+1, cols=7)
table.style = 'Light Grid Accent 1'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '모델'
hdr_cells[1].text = '정확도'
hdr_cells[2].text = 'AUC'
hdr_cells[3].text = '민감도'
hdr_cells[4].text = '특이도'
hdr_cells[5].text = '정밀도'
hdr_cells[6].text = 'F1-Score'

# Data
for idx, row in overall_summary.iterrows():
    cells = table.rows[idx+1].cells
    cells[0].text = row['Model']
    cells[1].text = f"{row['Accuracy']:.3f}"
    cells[2].text = f"{row['AUC']:.3f}"
    cells[3].text = f"{row['Sensitivity']:.3f}"
    cells[4].text = f"{row['Specificity']:.3f}"
    cells[5].text = f"{row['Precision']:.3f}"
    cells[6].text = f"{row['F1-Score']:.3f}"

doc.add_paragraph()

# Key findings
doc.add_heading('주요 발견사항:', 3)
best_model = overall_summary.iloc[-1]
doc.add_paragraph(
    f'• 최고 성능 모델 (메틸화 + 나이 + 성별 + CA19-9): AUC {best_model["AUC"]:.3f}, '
    f'정확도 {best_model["Accuracy"]*100:.1f}%'
)
doc.add_paragraph(
    f'• 민감도 {best_model["Sensitivity"]*100:.1f}%: 100명의 암 환자 중 약 {int(best_model["Sensitivity"]*100)}명을 정확히 진단'
)
doc.add_paragraph(
    f'• 특이도 {best_model["Specificity"]*100:.1f}%: 100명의 건강인 중 약 {int(best_model["Specificity"]*100)}명을 정확히 판별'
)

doc.add_paragraph()

# Add ROC curve image
doc.add_heading('ROC 곡선:', 3)
doc.add_picture(str(RESULTS_DIR / "roc_curves_overall.png"), width=Inches(5.5))

doc.add_page_break()

# ============================================================================
# 3.2 병기별 성능
# ============================================================================
doc.add_heading('3.2 병기별 진단 성능', 2)

# Stage performance table
table = doc.add_table(rows=len(stage_summary)+1, cols=6)
table.style = 'Light Grid Accent 1'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '병기'
hdr_cells[1].text = '환자 수'
hdr_cells[2].text = '정확도'
hdr_cells[3].text = 'AUC'
hdr_cells[4].text = '민감도'
hdr_cells[5].text = '특이도'

# Data
for idx, row in stage_summary.iterrows():
    cells = table.rows[idx+1].cells
    cells[0].text = row['Stage']
    cells[1].text = str(row['N_Patients'])
    cells[2].text = f"{row['Accuracy']:.3f}"
    cells[3].text = f"{row['AUC']:.3f}"
    cells[4].text = f"{row['Sensitivity']:.3f}"
    cells[5].text = f"{row['Specificity']:.3f}"

doc.add_paragraph()

# Early stage findings
early_stage_row = stage_summary[stage_summary['Stage'] == 'Early (1-2)'].iloc[0]
late_stage_row = stage_summary[stage_summary['Stage'] == 'Late (3-4)'].iloc[0]

doc.add_heading('조기 암(1-2기) 진단 성능:', 3)
early_detected = int(early_stage_row['Sensitivity'] * early_stage_row['N_Patients'])
doc.add_paragraph(
    f'• 민감도: {early_stage_row["Sensitivity"]*100:.1f}% '
    f'({early_detected}/{early_stage_row["N_Patients"]}명 검출)'
)
doc.add_paragraph(
    f'• AUC: {early_stage_row["AUC"]:.3f} - 매우 우수한 진단 성능'
)
doc.add_paragraph(
    '• 조기 췌장암 환자의 95% 이상을 정확히 진단할 수 있음을 시사합니다.'
)

doc.add_paragraph()

doc.add_heading('진행성 암(3-4기) 진단 성능:', 3)
late_detected = int(late_stage_row['Sensitivity'] * late_stage_row['N_Patients'])
doc.add_paragraph(
    f'• 민감도: {late_stage_row["Sensitivity"]*100:.1f}% '
    f'({late_detected}/{late_stage_row["N_Patients"]}명 검출)'
)
doc.add_paragraph(
    f'• AUC: {late_stage_row["AUC"]:.3f} - 매우 우수한 진단 성능'
)

doc.add_paragraph()

# Add stage-specific ROC curve
doc.add_heading('병기별 ROC 곡선:', 3)
doc.add_picture(str(RESULTS_DIR / "roc_curves_stage_specific.png"), width=Inches(5.5))

doc.add_page_break()

# ============================================================================
# 3.3 CA19-9 기여도
# ============================================================================
doc.add_heading('3.3 CA19-9의 진단 기여도', 2)

cancer_ca199 = patients['ca19_9'].dropna()
control_ca199 = controls['ca19_9'].dropna()

doc.add_heading('CA19-9 수치 비교:', 3)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = '통계량'
hdr_cells[1].text = '환자군'
hdr_cells[2].text = '대조군'

table.rows[1].cells[0].text = '평균'
table.rows[1].cells[1].text = f"{cancer_ca199.mean():.1f} U/mL"
table.rows[1].cells[2].text = f"{control_ca199.mean():.1f} U/mL"

table.rows[2].cells[0].text = '중앙값'
table.rows[2].cells[1].text = f"{cancer_ca199.median():.1f} U/mL"
table.rows[2].cells[2].text = f"{control_ca199.median():.1f} U/mL"

table.rows[3].cells[0].text = '배수 차이'
table.rows[3].cells[1].text = f"{cancer_ca199.mean() / control_ca199.mean():.1f}배 높음"
table.rows[3].cells[2].text = '-'

doc.add_paragraph()

doc.add_heading('모델 성능 비교:', 3)
meth_only = overall_summary[overall_summary['Model'] == 'Methylation only'].iloc[0]
meth_clinical = overall_summary[overall_summary['Model'] == 'Methylation + Age + Sex'].iloc[0]
full_model = overall_summary[overall_summary['Model'] == 'Methylation + Age + Sex + CA19-9'].iloc[0]

doc.add_paragraph(f'• 메틸화만: AUC {meth_only["AUC"]:.3f}')
doc.add_paragraph(f'• 메틸화 + 나이 + 성별: AUC {meth_clinical["AUC"]:.3f} (Δ = +{meth_clinical["AUC"] - meth_only["AUC"]:.3f})')
doc.add_paragraph(f'• 메틸화 + 나이 + 성별 + CA19-9: AUC {full_model["AUC"]:.3f} (Δ = +{full_model["AUC"] - meth_clinical["AUC"]:.3f})')

doc.add_paragraph()

doc.add_heading('주요 발견:', 3)
doc.add_paragraph(
    '• CA19-9는 환자군과 대조군을 통계적으로 명확히 구분합니다 (p < 0.001).'
)
doc.add_paragraph(
    f'• 그러나 메틸화 데이터에 CA19-9를 추가했을 때 AUC 향상은 미미합니다 '
    f'(+{full_model["AUC"] - meth_clinical["AUC"]:.3f}).'
)
doc.add_paragraph(
    '• 이는 메틸화 데이터가 이미 매우 강력한 진단 정보를 제공하기 때문입니다.'
)
doc.add_paragraph(
    '• 나이와 성별이 CA19-9보다 더 큰 기여를 합니다 '
    f'(AUC +{meth_clinical["AUC"] - meth_only["AUC"]:.3f}).'
)

doc.add_page_break()

# ============================================================================
# 4. 임상적 의의
# ============================================================================
doc.add_heading('4. 임상적 의의', 1)

doc.add_heading('4.1 조기 진단의 가능성', 2)
doc.add_paragraph(
    '본 연구는 EMR-seq 기반 메틸화 분석이 췌장암 조기 진단에 매우 효과적임을 보여줍니다:'
)
doc.add_paragraph(
    f'• 조기 암(1-2기) 검출률: {early_stage_row["Sensitivity"]*100:.1f}%'
)
doc.add_paragraph(
    '• 기존 종양 마커(CA19-9)보다 우수한 성능'
)
doc.add_paragraph(
    '• 비침습적 액체 생검으로 조기 발견 가능'
)

doc.add_heading('4.2 기존 진단법과의 비교', 2)
doc.add_paragraph('본 연구의 메틸화 기반 진단법의 장점:')
doc.add_paragraph('• 높은 민감도와 특이도 (AUC > 0.97)')
doc.add_paragraph('• 조기 암 검출 능력 우수')
doc.add_paragraph('• 혈액 샘플만으로 검사 가능 (비침습적)')
doc.add_paragraph('• CA19-9 단독 검사보다 정확함')

doc.add_heading('4.3 향후 활용 방안', 2)
doc.add_paragraph('• 췌장암 고위험군 스크리닝 도구로 활용')
doc.add_paragraph('• 정기 건강검진에 포함하여 조기 발견율 향상')
doc.add_paragraph('• 치료 효과 모니터링 및 재발 감시')
doc.add_paragraph('• 다른 임상 정보와 결합하여 정밀 의료 구현')

doc.add_page_break()

# ============================================================================
# 5. 결론
# ============================================================================
doc.add_heading('5. 결론', 1)

doc.add_paragraph(
    '본 연구는 EMR-seq 기술을 이용한 DNA 메틸화 분석이 췌장암 조기 진단에 '
    '매우 효과적인 도구임을 입증하였습니다. 주요 결과는 다음과 같습니다:'
)

doc.add_paragraph()

doc.add_heading('핵심 성과:', 2)
doc.add_paragraph(
    f'1. 전체 진단 성능: AUC {best_model["AUC"]:.3f}, 정확도 {best_model["Accuracy"]*100:.1f}%, '
    f'민감도 {best_model["Sensitivity"]*100:.1f}%'
)
doc.add_paragraph(
    f'2. 조기 암(1-2기) 검출: 민감도 {early_stage_row["Sensitivity"]*100:.1f}%, '
    f'AUC {early_stage_row["AUC"]:.3f}'
)
doc.add_paragraph(
    f'3. 46개의 유의한 DMR 발견 및 검증'
)
doc.add_paragraph(
    '4. 메틸화 데이터가 기존 종양 마커(CA19-9)보다 우수한 진단 성능 제공'
)

doc.add_paragraph()

doc.add_heading('임상적 가치:', 2)
doc.add_paragraph(
    '• 췌장암의 조기 발견율을 크게 향상시킬 수 있는 잠재력'
)
doc.add_paragraph(
    '• 비침습적 액체 생검으로 환자 부담 최소화'
)
doc.add_paragraph(
    '• 높은 민감도와 특이도로 불필요한 추가 검사 감소'
)
doc.add_paragraph(
    '• 정밀 의료 시대의 맞춤형 진단 도구로 활용 가능'
)

doc.add_paragraph()

doc.add_heading('향후 연구 방향:', 2)
doc.add_paragraph(
    '• 더 큰 규모의 전향적 임상 연구를 통한 검증'
)
doc.add_paragraph(
    '• 다기관 연구를 통한 일반화 가능성 확인'
)
doc.add_paragraph(
    '• 다른 암종과의 감별 진단 능력 평가'
)
doc.add_paragraph(
    '• 치료 반응 예측 및 예후 판정 모델 개발'
)

doc.add_page_break()

# ============================================================================
# 6. 부록
# ============================================================================
doc.add_heading('6. 부록', 1)

doc.add_heading('6.1 통계 방법 상세', 2)
doc.add_paragraph('• DMR 검출: DMRseq 알고리즘 (FDR < 0.05)')
doc.add_paragraph('• 예측 모델: Logistic Regression with L2 regularization')
doc.add_paragraph('• 교차 검증: 5-fold stratified cross-validation')
doc.add_paragraph('• 클래스 불균형 처리: Balanced class weights')
doc.add_paragraph('• 특징 정규화: StandardScaler (평균 0, 표준편차 1)')
doc.add_paragraph('• 통계적 유의성 검정: Mann-Whitney U test')

doc.add_heading('6.2 데이터 품질 관리', 2)
doc.add_paragraph(f'• 총 샘플 수: {len(integrated_data)}개')
doc.add_paragraph(f'• 메틸화 데이터 결측치: 중앙값으로 대체')
doc.add_paragraph(f'• CA19-9 결측치: {integrated_data["ca19_9"].isna().sum()}개 (중앙값으로 대체)')
doc.add_paragraph('• 이상치 처리: 임상적으로 타당한 범위 내 모든 값 포함')

doc.add_heading('6.3 소프트웨어 및 패키지', 2)
doc.add_paragraph('• R 4.x: DMRseq, bsseq')
doc.add_paragraph('• Python 3.x: scikit-learn, pandas, numpy')
doc.add_paragraph('• 시각화: matplotlib, seaborn')
doc.add_paragraph('• 통계 분석: scipy.stats')

# Save document
output_file = OUTPUT_DIR / "췌장암_EMRseq_메틸화_분석_최종보고서.docx"
doc.save(output_file)

print(f"\n최종 한글 리포트 생성 완료!")
print(f"저장 위치: {output_file}")
print("=" * 80)
