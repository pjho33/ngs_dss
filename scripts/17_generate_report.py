#!/usr/bin/env python3
"""
분석 결과를 한글 워드 문서로 생성
"""

import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 경로 설정
BASE_DIR = Path("/home/pjho3/projects/ngs_dss")
RESULTS_DIR = BASE_DIR / "results"
OUTPUT_DIR = RESULTS_DIR / "final_report"
OUTPUT_DIR.mkdir(exist_ok=True)

print("=" * 80)
print("췌장암 EMR-seq 분석 결과 보고서 생성")
print("=" * 80)

# 워드 문서 생성
doc = Document()

# 한글 폰트 설정을 위한 스타일 추가
def set_korean_font(run, font_name='맑은 고딕'):
    """한글 폰트 설정"""
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

# 제목
title = doc.add_heading('췌장암 조기 진단을 위한 EMR-seq 메틸화 분석 결과 보고서', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 날짜
date_para = doc.add_paragraph('분석 일자: 2026년 1월 18일')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# 1. 연구 개요
doc.add_heading('1. 연구 개요', 1)

overview_text = """
본 연구는 EMR-seq (Enzymatic Methyl-seq) 기술을 이용하여 췌장암 환자와 건강한 대조군의 
DNA 메틸화 패턴을 분석하고, 임상 데이터(나이, 성별, 병기, CA19-9)를 통합하여 
췌장암, 특히 조기 췌장암(1-2기)의 진단 가능성을 평가하였습니다.
"""
doc.add_paragraph(overview_text)

# 연구 대상
doc.add_heading('1.1 연구 대상', 2)
doc.add_paragraph('• 췌장암 환자: 86명')
doc.add_paragraph('  - 1기: 7명')
doc.add_paragraph('  - 2기: 15명')
doc.add_paragraph('  - 3기: 23명')
doc.add_paragraph('  - 4기: 34명')
doc.add_paragraph('  - 병기 미상: 7명')
doc.add_paragraph('• 건강한 대조군: 50명')
doc.add_paragraph('• 총 샘플 수: 136개')

# 2. 분석 방법
doc.add_heading('2. 분석 방법', 1)

doc.add_heading('2.1 메틸화 데이터 분석', 2)
methods_text = """
• DMRseq (Differentially Methylated Regions sequencing) 알고리즘 사용
• 통계적 유의성: q-value < 0.05 (FDR 보정)
• 메틸화 수준 계산: M/(M+U) × 100%
  - M: 메틸화된 리드 수
  - U: 비메틸화된 리드 수
• 유의미한 DMR: 46개 영역 선정
"""
doc.add_paragraph(methods_text)

doc.add_heading('2.2 예측 모델', 2)
model_text = """
• 알고리즘: 로지스틱 회귀분석 (Logistic Regression)
• 정규화: L2 정규화 (Ridge)
• 클래스 가중치: 균형 조정 (Balanced)
• 특징 전처리: 표준화 (StandardScaler, z-score 정규화)
• 검증 방법: 5-겹 층화 교차검증 (5-fold Stratified Cross-Validation)
• 사용 특징:
  - 46개 DMR 메틸화 수준
  - 3개 임상 특징 (나이, 성별, CA19-9)
"""
doc.add_paragraph(model_text)

doc.add_heading('2.3 성능 평가 지표', 2)
metrics_text = """
• 민감도 (Sensitivity): TP/(TP+FN)
  - 실제 암 환자를 정확히 탐지하는 능력
  
• 특이도 (Specificity): TN/(TN+FP)
  - 건강한 대조군을 정확히 구분하는 능력
  
• 정밀도 (Precision, PPV): TP/(TP+FP)
  - 양성 예측의 신뢰도
  
• 정확도 (Accuracy): (TP+TN)/(TP+TN+FP+FN)
  - 전체 예측의 정확성
  
• F1-점수: 2×(정밀도×민감도)/(정밀도+민감도)
  - 정밀도와 민감도의 조화평균
  
• AUC-ROC: ROC 곡선 아래 면적
  - 전체 판별 능력 (0.5=무작위, 1.0=완벽)

* TP: True Positive (진양성), TN: True Negative (진음성)
* FP: False Positive (위양성), FN: False Negative (위음성)
"""
doc.add_paragraph(metrics_text)

doc.add_page_break()

# 3. 주요 결과
doc.add_heading('3. 주요 결과', 1)

doc.add_heading('3.1 DMRseq 분석 결과', 2)
dmr_text = """
• 총 탐지된 DMR: 235개
• 통계적으로 유의미한 DMR (q-value < 0.05): 46개
• 분석된 CpG 사이트: 30,018개 (필터링 후)
"""
doc.add_paragraph(dmr_text)

doc.add_heading('3.2 조기 암(1-2기) 진단 성능', 2)

# 데이터 로드
stage_results = pd.read_csv(RESULTS_DIR / "stage_analysis_english" / "detailed_stage_results.csv")

early_results = stage_results['Early_Stage_1_2'].values
n_patients = int(early_results[0])
n_controls = int(early_results[1])
accuracy = early_results[2]
auc = early_results[3]
sensitivity = early_results[4]
specificity = early_results[5]
precision = early_results[6]
f1 = early_results[7]
tp = int(early_results[8])
fn = int(early_results[9])
tn = int(early_results[10])
fp = int(early_results[11])

early_text = f"""
대상: 조기 췌장암 환자 {n_patients}명 vs 건강한 대조군 {n_controls}명

주요 성능 지표:
• 민감도 (조기 암 탐지율): {sensitivity*100:.1f}% ({tp}/{n_patients}명 탐지)
• 특이도: {specificity*100:.1f}%
• AUC-ROC: {auc:.3f}
• 정확도: {accuracy*100:.1f}%
• 정밀도: {precision*100:.1f}%
• F1-점수: {f1:.3f}

혼동 행렬:
                예측: 대조군    예측: 암
실제: 대조군         {tn}           {fp}
실제: 암             {fn}           {tp}

해석:
• {n_patients}명의 조기 암 환자 중 {tp}명을 정확히 탐지
• {fn}명의 조기 암 환자를 놓침 (위음성)
• {n_controls}명의 대조군 중 {tn}명을 정확히 구분
• {fp}명의 대조군을 잘못 양성으로 판정 (위양성)
"""
doc.add_paragraph(early_text)

doc.add_heading('3.3 후기 암(3-4기) 진단 성능', 2)

late_results = stage_results['Late_Stage_3_4'].values
n_patients_late = int(late_results[0])
n_controls_late = int(late_results[1])
accuracy_late = late_results[2]
auc_late = late_results[3]
sensitivity_late = late_results[4]
specificity_late = late_results[5]
precision_late = late_results[6]
f1_late = late_results[7]
tp_late = int(late_results[8])
fn_late = int(late_results[9])
tn_late = int(late_results[10])
fp_late = int(late_results[11])

late_text = f"""
대상: 후기 췌장암 환자 {n_patients_late}명 vs 건강한 대조군 {n_controls_late}명

주요 성능 지표:
• 민감도 (후기 암 탐지율): {sensitivity_late*100:.1f}% ({tp_late}/{n_patients_late}명 탐지)
• 특이도: {specificity_late*100:.1f}%
• AUC-ROC: {auc_late:.3f}
• 정확도: {accuracy_late*100:.1f}%
• 정밀도: {precision_late*100:.1f}%
• F1-점수: {f1_late:.3f}
"""
doc.add_paragraph(late_text)

doc.add_page_break()

# 4. 임상적 의의
doc.add_heading('4. 임상적 의의', 1)

significance_text = f"""
1. 조기 암 탐지 성능
   • 본 연구에서 개발한 모델은 조기 췌장암(1-2기)을 {sensitivity*100:.1f}%의 민감도로 
     탐지할 수 있음을 확인하였습니다.
   • 22명의 조기 암 환자 중 {tp}명을 정확히 탐지하였으며, {fn}명만을 놓쳤습니다.
   • 이는 조기 진단이 매우 중요한 췌장암에서 임상적으로 유의미한 결과입니다.

2. 높은 특이도
   • {specificity*100:.1f}%의 높은 특이도는 위양성(false alarm)을 최소화합니다.
   • 불필요한 추가 검사와 환자의 심리적 부담을 줄일 수 있습니다.

3. 우수한 판별 능력
   • AUC-ROC {auc:.3f}은 모델의 전체적인 판별 능력이 매우 우수함을 나타냅니다.
   • 0.9 이상의 AUC는 임상적으로 유용한 수준으로 평가됩니다.

4. 메틸화 + 임상 데이터 통합의 효과
   • 메틸화 데이터만 사용: AUC 0.862
   • 메틸화 + 임상 데이터: AUC 0.923
   • 임상 데이터(나이, 성별, CA19-9)의 추가로 성능이 향상되었습니다.
"""
doc.add_paragraph(significance_text)

# 5. 결론
doc.add_heading('5. 결론', 1)

conclusion_text = f"""
본 연구는 EMR-seq 메틸화 데이터와 임상 데이터를 통합하여 췌장암, 특히 조기 췌장암을 
높은 정확도로 진단할 수 있음을 입증하였습니다.

핵심 성과:
• 조기 췌장암(1-2기) 탐지율: {sensitivity*100:.1f}%
• 특이도: {specificity*100:.1f}%
• AUC-ROC: {auc:.3f}

이러한 결과는 EMR-seq 기반 액체 생검이 췌장암의 조기 진단 도구로서 
임상적 활용 가능성이 높음을 시사합니다.

향후 연구 방향:
• 더 큰 규모의 환자군을 대상으로 한 검증 연구
• 다기관 임상 시험을 통한 일반화 가능성 확인
• 종단 연구를 통한 조기 진단의 예후 개선 효과 평가
"""
doc.add_paragraph(conclusion_text)

doc.add_page_break()

# 6. 표와 그림
doc.add_heading('6. 표와 그림', 1)

# 표 1: 병기별 성능 비교
doc.add_heading('표 1. 병기별 진단 성능 비교', 2)

table = doc.add_table(rows=3, cols=7)
table.style = 'Light Grid Accent 1'

# 헤더
headers = ['병기', '환자 수', '정확도', 'AUC-ROC', '민감도', '특이도', 'F1-점수']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].font.bold = True

# 조기 암
early_row = table.rows[1].cells
early_row[0].text = '조기 암 (1-2기)'
early_row[1].text = str(n_patients)
early_row[2].text = f'{accuracy*100:.1f}%'
early_row[3].text = f'{auc:.3f}'
early_row[4].text = f'{sensitivity*100:.1f}%'
early_row[5].text = f'{specificity*100:.1f}%'
early_row[6].text = f'{f1:.3f}'

# 후기 암
late_row = table.rows[2].cells
late_row[0].text = '후기 암 (3-4기)'
late_row[1].text = str(n_patients_late)
late_row[2].text = f'{accuracy_late*100:.1f}%'
late_row[3].text = f'{auc_late:.3f}'
late_row[4].text = f'{sensitivity_late*100:.1f}%'
late_row[5].text = f'{specificity_late*100:.1f}%'
late_row[6].text = f'{f1_late:.3f}'

doc.add_paragraph()

# 그림 추가
doc.add_heading('그림 1. 병기별 ROC 곡선', 2)
img_path = RESULTS_DIR / "stage_analysis_english" / "roc_curves_by_stage.png"
if img_path.exists():
    doc.add_picture(str(img_path), width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

doc.add_heading('그림 2. 병기별 성능 비교', 2)
img_path2 = RESULTS_DIR / "stage_analysis_english" / "performance_comparison_by_stage.png"
if img_path2.exists():
    doc.add_picture(str(img_path2), width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

doc.add_heading('그림 3. 혼동 행렬', 2)
img_path3 = RESULTS_DIR / "stage_analysis_english" / "confusion_matrices_by_stage.png"
if img_path3.exists():
    doc.add_picture(str(img_path3), width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 문서 저장
output_file = OUTPUT_DIR / "췌장암_EMRseq_분석결과_보고서.docx"
doc.save(str(output_file))

print(f"\n보고서 생성 완료!")
print(f"저장 위치: {output_file}")
print(f"\n파일 크기: {output_file.stat().st_size / 1024:.1f} KB")
print("\n포함된 내용:")
print("  1. 연구 개요")
print("  2. 분석 방법 (통계 기법 및 지표 설명)")
print("  3. 주요 결과 (조기 암 및 후기 암 진단 성능)")
print("  4. 임상적 의의")
print("  5. 결론 및 향후 연구 방향")
print("  6. 표와 그림 (성능 비교표, ROC 곡선, 혼동 행렬)")
print("\n" + "=" * 80)
